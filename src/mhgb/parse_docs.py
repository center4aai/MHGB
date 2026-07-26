"""
Шаг 1: Парсинг docx-файлов НПА → структурированный JSONL.

Формат записи:
{
  "id": "ТК_261",
  "law": "Трудовой кодекс Российской Федерации",
  "law_short": "ТК",
  "article": "261",
  "title": "Гарантии беременной женщине...",
  "paragraphs": ["текст абзаца 1", "текст абзаца 2", ...],
  "text": "полный текст статьи",
  "valid_from": "2024-02-25",   # дата последней редакции, если известна
  "valid_to": null,
  "version": "ред. от 25.02.2024"
}
"""

import argparse
import hashlib
import json
import re
import sys
from pathlib import Path
from docx import Document

from mhgb.logging_setup import logger
from mhgb.utils.branch_mapping import BRANCH_OF_LAW

# Русские названия месяцев → номер
MONTHS = {
    "января": 1, "февраля": 2, "марта": 3, "апреля": 4,
    "мая": 5, "июня": 6, "июля": 7, "августа": 8,
    "сентября": 9, "октября": 10, "ноября": 11, "декабря": 12,
}

# Имя файла → короткое название закона (все 10 НПА)
LAW_SHORT = {
    "1. Конституция РФ.docx":           "КРФ",
    "8. Гражданский кодекс РФ.docx":    "ГК",
    "9. Трудовой кодекс РФ.docx":       "ТК",
    "10. КоАП РФ.docx":                 "КоАП",
    "11. Уголовный кодекс РФ.docx":     "УК",
    "12. Семейный кодекс РФ.docx":      "СК",
    "13. Налоговый кодекс РФ.docx":     "НК",
    "14. Земельный кодекс РФ.docx":     "ЗК",
    "15. Жилищный кодекс РФ.docx":      "ЖК",
    "16. Градостроительный кодекс РФ.docx": "ГрК",
}

# Стили, которые игнорируем (по умолчанию)
SKIP_STYLES = {
    "Комментарий",
    "Информация о версии",
    "Информация об изменениях",
    "Подзаголовок для информации об изменениях",
    "Прижатый влево",
    "Таблицы (моноширинный)",
}

# Переопределения стилей для docx с нестандартной разметкой
# header  — стиль заголовка статьи
# body    — стиль текста абзаца
# version — стиль строк версии/редакции (None = не используется)
STYLE_OVERRIDES: dict[str, dict[str, str | None]] = {
    "КРФ": {
        "header":  "t",
        "body":    "Normal (Web)",
        "version": None,
    },
}


def parse_russian_date(text: str) -> str | None:
    """Извлекает первую дату вида '25 февраля 2024 г.' из текста, возвращает 'YYYY-MM-DD'."""
    m = re.search(r"(\d{1,2})\s+(января|февраля|марта|апреля|мая|июня|июля|августа|сентября|октября|ноября|декабря)\s+(\d{4})", text)
    if m:
        day, month_name, year = int(m.group(1)), m.group(2), int(m.group(3))
        return f"{year:04d}-{MONTHS[month_name]:02d}-{day:02d}"
    return None


def parse_article_header(text: str) -> tuple[str, str] | None:
    """
    Из строки вида 'Статья 261. Гарантии беременной женщине...'
    возвращает ('261', 'Гарантии беременной женщине...').
    Поддерживает также 'Статья 261.1.' и просто 'Статья 261' без названия.
    """
    m = re.match(r"Стать[яи]\s+([\d.]+)\.?\s*(.*)", text, re.DOTALL)
    if m:
        num = m.group(1).rstrip(".")
        title = m.group(2).strip()
        return num, title
    return None


def extract_law_name(doc: Document) -> str:
    """Берёт первый непустой Heading 1 как полное название закона."""
    for p in doc.paragraphs:
        if p.style.name == "Heading 1" and p.text.strip():
            return p.text.strip()
    return "Неизвестный закон"


def parse_docx(path: Path, law_short: str) -> list[dict]:
    """Парсит один docx, возвращает список статей."""
    doc = Document(path)
    law_name = extract_law_name(doc)

    styles = {
        "header": "Заголовок статьи",
        "body": "Normal",
        "version": "Информация о версии",
        **STYLE_OVERRIDES.get(law_short, {}),
    }
    HEADER_STYLE  = styles["header"]
    BODY_STYLE    = styles["body"]
    VERSION_STYLE = styles["version"]

    articles = []
    current_article_num = None
    current_title = ""
    current_paragraphs = []
    current_article_version_texts = []  # версионные тексты текущей статьи
    pending_version_texts = []          # версионные тексты, накопленные ДО следующего заголовка

    def flush(article_num, title, paragraphs, version_texts):
        if article_num is None:
            return

        # Определяем, утратила ли статья силу
        title_lower = title.lower()
        is_repealed = bool(re.search(r"утратил[аио]?\s+силу", title_lower))
        valid_to = parse_russian_date(title) if is_repealed else None

        # Полностью утратившие силу статьи сохраняем даже без тела
        if not paragraphs and not is_repealed:
            return

        valid_from = None
        version_str = None
        for vt in version_texts:
            d = parse_russian_date(vt)
            if d:
                valid_from = d
                m = re.search(r"(\d{1,2}\s+\w+\s+\d{4})", vt)
                if m:
                    version_str = "ред. от " + m.group(1)
                break

        # Частичные утраты силы: ищем в тексте абзацев
        has_repealed_parts = any(
            re.search(r"утратил[аио]?\s+силу", p.lower())
            for p in paragraphs
        )

        text = "\n".join(p for p in paragraphs if p)
        articles.append({
            "id": f"{law_short}_{article_num}",
            "law": law_name,
            "law_short": law_short,
            "branch_of_law": BRANCH_OF_LAW.get(law_short, "неизвестное"),
            "article": article_num,
            "title": title,
            "paragraphs": paragraphs,
            "text": text,
            "valid_from": valid_from,
            "valid_to": valid_to,
            "is_repealed": is_repealed,
            "has_repealed_parts": has_repealed_parts,
            "version": version_str,
        })

    for p in doc.paragraphs:
        text = p.text.strip()
        style = p.style.name

        if not text:
            continue

        if style in SKIP_STYLES:
            # Версионные тексты накапливаем — они принадлежат СЛЕДУЮЩЕЙ статье
            if VERSION_STYLE and style == VERSION_STYLE:
                pending_version_texts.append(text)
            continue

        if style == HEADER_STYLE:
            parsed = parse_article_header(text)
            if not parsed:
                continue  # заголовок раздела/части — не статья
            # Сохраняем предыдущую статью с ЕЁ версионными текстами
            flush(current_article_num, current_title, current_paragraphs, current_article_version_texts)
            current_article_num, current_title = parsed
            current_paragraphs = []
            current_article_version_texts = pending_version_texts
            pending_version_texts = []
            continue

        if style == "Heading 1":
            continue

        if style == BODY_STYLE:
            if current_article_num is not None:
                current_paragraphs.append(text)
            continue

    # Последняя статья
    flush(current_article_num, current_title, current_paragraphs, current_article_version_texts)
    return articles


def main():
    parser = argparse.ArgumentParser(description="MHGB: парсинг НПА → corpus.jsonl")
    parser.add_argument(
        "--codes", nargs="*",
        help="Коды НПА для парсинга (напр. --codes ТК ГК). Без аргумента — все.",
    )
    parser.add_argument(
        "--tag", default="v2_full_corpus",
        help="Тег версии корпуса для MongoDB (default: v2_full_corpus)",
    )
    args = parser.parse_args()

    from mhgb.config import cfg
    cfg.data_dir.mkdir(exist_ok=True)

    # Фильтр по --codes
    law_map = dict(LAW_SHORT)
    if args.codes:
        codes_set = set(args.codes)
        unknown = codes_set - set(LAW_SHORT.values())
        if unknown:
            logger.warning("Неизвестные коды НПА: {}", ", ".join(sorted(unknown)))
        law_map = {f: s for f, s in LAW_SHORT.items() if s in codes_set}

    all_articles = []
    for fname, short in law_map.items():
        path = cfg.legal_docs_dir / fname
        if not path.exists():
            logger.warning("[SKIP] не найден: {}", fname)
            continue
        articles = parse_docx(path, short)
        n_repealed = sum(1 for a in articles if a["is_repealed"])
        n_with_date = sum(1 for a in articles if a["valid_from"])
        logger.info("{:6s}  {:4d} статей  ({} утр. силу, {} с датой ред.)",
                    short, len(articles), n_repealed, n_with_date)
        all_articles.extend(articles)

    if not all_articles:
        logger.error("Нет статей для сохранения. Проверьте наличие docx в {}", cfg.legal_docs_dir)
        sys.exit(1)

    # Сохраняем corpus.jsonl
    corpus_path = cfg.data_dir / "corpus.jsonl"
    with open(corpus_path, "w", encoding="utf-8") as f:
        for a in all_articles:
            f.write(json.dumps(a, ensure_ascii=False) + "\n")

    with_date = sum(1 for a in all_articles if a["valid_from"])
    logger.info("Итого: {} статей → {}", len(all_articles), corpus_path)
    logger.info("С датой редакции: {} ({}%)", with_date, 100 * with_date // len(all_articles))

    branch_counts: dict[str, int] = {}
    for a in all_articles:
        b = a["branch_of_law"]
        branch_counts[b] = branch_counts.get(b, 0) + 1
    logger.info("Распределение по отраслям права:")
    for branch, cnt in sorted(branch_counts.items(), key=lambda x: -x[1]):
        logger.info("  {:25s} {:4d} статей", branch, cnt)

    # MongoDB версионирование
    checksum = hashlib.sha256(corpus_path.read_bytes()).hexdigest()
    law_counts = {}
    for a in all_articles:
        law_counts[a["law_short"]] = law_counts.get(a["law_short"], 0) + 1

    try:
        from mhgb.storage.mongo_client import MongoStorage
        from mhgb.storage.schemas import CorpusVersion
        from pymongo.errors import DuplicateKeyError

        storage = MongoStorage()
        if storage.check_connection():
            storage.ensure_collections()
            version = CorpusVersion(
                tag=args.tag,
                article_count=len(all_articles),
                law_counts=law_counts,
                checksum=checksum,
            )
            try:
                oid = storage.insert("corpus_versions", version.model_dump())
                logger.info("MongoDB corpus_versions: тег='{}', _id={}", args.tag, oid)
            except DuplicateKeyError:
                logger.warning(
                    "MongoDB: тег '{}' уже существует — пропущено. "
                    "Используйте --tag для другого имени.", args.tag
                )
        else:
            logger.warning("MongoDB недоступна — версионирование пропущено")
        storage.close()
    except Exception as e:
        logger.warning("MongoDB versioning error: {}", e)


if __name__ == "__main__":
    sys.stdout.reconfigure(encoding="utf-8")
    main()
